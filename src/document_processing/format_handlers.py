"""
format_handlers.py — Multi-format document text extraction dispatcher.

Supports:
  - PDF  (.pdf)   — via pdf_parser.py (pymupdf4llm → pymupdf fallback → OCR fallback)
  - DOCX (.docx)  — via python-docx
  - Images (.png, .jpg, .jpeg, .tiff, .tif, .bmp) — via Tesseract OCR

All handlers return a common schema: List[{"page_number": int, "markdown": str}]
"""

import os
import logging
from typing import List, Dict

logger = logging.getLogger(__name__)

# Extensions mapped to handler functions
_PDF_EXTENSIONS = {".pdf"}
_DOCX_EXTENSIONS = {".docx"}
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"}

SUPPORTED_EXTENSIONS = _PDF_EXTENSIONS | _DOCX_EXTENSIONS | _IMAGE_EXTENSIONS


def extract_text(file_path: str) -> List[Dict]:
    """Extract text from a document file, dispatching by extension.

    Args:
        file_path: Path to the document file.

    Returns:
        List of {"page_number": int, "markdown": str} for each page/section.

    Raises:
        ValueError: If the file extension is not supported.
        FileNotFoundError: If the file does not exist.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext in _PDF_EXTENSIONS:
        return _extract_pdf(file_path)
    elif ext in _DOCX_EXTENSIONS:
        return _extract_docx(file_path)
    elif ext in _IMAGE_EXTENSIONS:
        return _extract_image(file_path)
    else:
        raise ValueError(
            f"Unsupported file format '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )


def _extract_pdf(file_path: str) -> List[Dict]:
    """Extract text from PDF with OCR fallback for scanned documents."""
    from document_processing.pdf_parser import parse_pdf_to_markdown
    from document_processing.ocr_fallback import needs_ocr, ocr_pdf

    pages = parse_pdf_to_markdown(file_path)

    # If native extraction is insufficient, try OCR
    if needs_ocr(pages):
        logger.info("[FormatHandler] Native PDF extraction insufficient, attempting OCR...")
        ocr_pages = ocr_pdf(file_path)
        if ocr_pages:
            # Use OCR results if they're better than native
            ocr_chars = sum(len(p.get("markdown", "")) for p in ocr_pages)
            native_chars = sum(len(p.get("markdown", "")) for p in pages)
            if ocr_chars > native_chars:
                logger.info(
                    "[FormatHandler] OCR produced better results (%d vs %d chars)",
                    ocr_chars, native_chars,
                )
                return ocr_pages
            else:
                logger.info("[FormatHandler] Keeping native extraction (OCR not better)")

    return pages


def _extract_docx(file_path: str) -> List[Dict]:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        logger.error(
            "[FormatHandler] python-docx is required for DOCX support. "
            "Install with: pip install python-docx"
        )
        return []

    logger.info("[FormatHandler] Extracting DOCX: %s", os.path.basename(file_path))

    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        # Detect heading styles and convert to markdown
        if para.style and para.style.name:
            style = para.style.name.lower()
            if "heading 1" in style:
                lines.append(f"# {text}")
            elif "heading 2" in style:
                lines.append(f"## {text}")
            elif "heading 3" in style:
                lines.append(f"### {text}")
            elif "heading 4" in style:
                lines.append(f"#### {text}")
            elif "title" in style:
                lines.append(f"# {text}")
            else:
                lines.append(text)
        else:
            lines.append(text)

    # Extract tables
    for table in doc.tables:
        lines.append("")  # blank line before table
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                # Add markdown table separator after header row
                lines.append("| " + " | ".join("---" for _ in cells) + " |")
        lines.append("")  # blank line after table

    full_text = "\n".join(lines).strip()

    if not full_text:
        return []

    # DOCX doesn't have page numbers, treat as single page
    return [{"page_number": 1, "markdown": full_text}]


def _extract_image(file_path: str) -> List[Dict]:
    """Extract text from an image file using Tesseract OCR."""
    from document_processing.ocr_fallback import ocr_image

    text = ocr_image(file_path)
    if not text:
        return []

    return [{"page_number": 1, "markdown": text}]
